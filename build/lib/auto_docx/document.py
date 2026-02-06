# """
# Word document generation module.

# Handles the creation of formatted Word documents with script execution results.
# """

# import re
# from pathlib import Path
# from typing import Optional, Union
# import shutil

# from docx import Document
# from docx.shared import Pt, RGBColor, Inches
# from docx.enum.text import WD_ALIGN_PARAGRAPH
# from docx.enum.style import WD_STYLE_TYPE

# from .executor import ExecutionResult

# # Default roll number - change this for perpetual use
# DEFAULT_ROLL_NO = "123102115"

# # Regex to remove invalid XML characters
# _INVALID_XML_CHARS = re.compile(
#     r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]'
# )


# def _sanitize_text(text: str) -> str:
#     """Remove characters that are invalid in XML."""
#     return _INVALID_XML_CHARS.sub('', text)


# class DocumentGenerator:
#     """Generates Word documents from script execution results."""
    
#     # Formatting constants
#     MONOSPACE_FONT = "Courier New"
#     CODE_FONT_SIZE = Pt(10)
#     HEADING_FONT_SIZE = Pt(14)
#     BODY_FONT_SIZE = Pt(11)
    
#     def __init__(self, include_source: bool = False, verbose: bool = False, roll_no: Optional[str] = None):
#         """
#         Initialize the document generator.
        
#         Args:
#             include_source: Whether to include script source code in the document
#             verbose: Enable verbose output to console
#             roll_no: Roll number to include in the document
#         """
#         self.include_source = include_source
#         self.verbose = verbose
#         self.roll_no = roll_no if roll_no else DEFAULT_ROLL_NO
    
#     def generate(
#         self,
#         result: ExecutionResult,
#         output_path: Optional[Union[str, Path]] = None,
#     ) -> Path:
#         """
#         Generate a Word document from execution results.
        
#         Args:
#             result: The execution result to document
#             output_path: Custom output path (optional, defaults to script_output.docx)
            
#         Returns:
#             Path to the generated document
#         """
#         # Determine output path
#         if output_path:
#             output_path = Path(output_path)
#         else:
#             filename = f"{self.roll_no}_{result.script_path.stem}.docx"
#             output_path = result.script_path.parent / filename
        
#         # Ensure parent directory exists
#         output_path.parent.mkdir(parents=True, exist_ok=True)
        
#         if self.verbose:
#             print(f"[INFO] Generating document: {output_path}")
        
#         # Create document
#         doc = Document()
#         self._setup_styles(doc)
        
#         # Add content sections
#         self._add_header(doc, result)
        
#         if self.include_source:
#             self._add_source_code(doc, result)
        
#         self._add_output_section(doc, result)

#         # Save document
#         doc.save(str(output_path))
        
#         # Clean up images directory after embedding
#         self._cleanup_images(result)
        
#         if self.verbose:
#             print(f"[INFO] Document saved successfully: {output_path}")
        
#         return output_path
    
#     def _cleanup_images(self, result: ExecutionResult) -> None:
#         """Remove the persistent images directory after document generation if desired."""
#         # Note: You might want to keep images if debugging. 
#         # For now, we only clean up if it's explicitly temporary.
#         pass
    
#     def _setup_styles(self, doc: Document) -> None:
#         """Set up custom styles for the document."""
#         styles = doc.styles
        
#         # Code style for monospaced text
#         if "Code" not in [s.name for s in styles]:
#             code_style = styles.add_style("Code", WD_STYLE_TYPE.PARAGRAPH)
#             code_style.font.name = self.MONOSPACE_FONT
#             code_style.font.size = self.CODE_FONT_SIZE
#             code_style.paragraph_format.space_before = Pt(6)
#             code_style.paragraph_format.space_after = Pt(6)
    
#     def _add_header(self, doc: Document, result: ExecutionResult) -> None:
#         """Add document header with filename as title and roll number."""
#         # Title - use the script filename (without extension)
#         title = doc.add_heading(result.script_path.stem, level=0)
#         title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
#         # Roll Number
#         if self.roll_no:
#             roll_para = doc.add_paragraph(f"Roll No: {self.roll_no}")
#             roll_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
#         doc.add_paragraph()  # Spacer
    
#     def _add_source_code(self, doc: Document, result: ExecutionResult) -> None:
#         """Add source code section."""
#         doc.add_heading("Source Code", level=1)
        
#         # Add source code with monospace formatting
#         if result.source_code and result.source_code.strip():
#             self._add_code_block(doc, result.source_code)
#         else:
#             doc.add_paragraph("(Empty file)")
        
#         doc.add_paragraph()  # Spacer
    
#     def _add_output_section(self, doc: Document, result: ExecutionResult) -> None:
#         """Add stdout output section."""
#         doc.add_heading("Output", level=1)
        
#         if result.output_items:
#             self._add_output_items(doc, result.output_items)
#         elif result.stdout.strip():
#             self._add_code_block(doc, result.stdout)
#         else:
#             doc.add_paragraph("(No output)")
        
#         doc.add_paragraph()  # Spacer
    
#     def _add_code_block(
#         self,
#         doc: Document,
#         code: str,
#         is_error: bool = False,
#     ) -> None:
#         """
#         Add a code block with monospace formatting.
#         """
#         # Sanitize and split into lines
#         code = _sanitize_text(code)
#         lines = code.split("\n")
        
#         for line in lines:
#             para = doc.add_paragraph()
#             run = para.add_run(line if line else " ")
#             run.font.name = self.MONOSPACE_FONT
#             run.font.size = self.CODE_FONT_SIZE
            
#             if is_error:
#                 run.font.color.rgb = RGBColor(180, 0, 0)
            
#             # Reduce paragraph spacing for code blocks
#             para.paragraph_format.space_before = Pt(0)
#             para.paragraph_format.space_after = Pt(0)
#             para.paragraph_format.line_spacing = 1.0

#     def _add_output_items(self, doc: Document, items: list) -> None:
#         """Add ordered output items (text, markdown, and images)."""
#         for item in items:
#             kind = getattr(item, "kind", None)
#             if kind == "image":
#                 self._add_image(doc, item.content)
#             elif kind == "markdown":
#                 self._add_markdown_block(doc, item.content)
#             else:
#                 text = item.content if hasattr(item, "content") else ""
#                 if text:
#                     self._add_code_block(doc, text)

#     def _add_markdown_block(self, doc: Document, content: str) -> None:
#         """Add markdown content as formatted text."""
#         content = _sanitize_text(content)
#         lines = content.strip().split("\n")
#         for line in lines:
#             if line.startswith("# "):
#                 doc.add_heading(_sanitize_text(line[2:]), level=1)
#             elif line.startswith("## "):
#                 doc.add_heading(_sanitize_text(line[3:]), level=2)
#             elif line.startswith("### "):
#                 doc.add_heading(_sanitize_text(line[4:]), level=3)
#             elif line.strip():
#                 doc.add_paragraph(_sanitize_text(line))
#         doc.add_paragraph()  # Spacer

#     def _add_image(self, doc: Document, image_path: str) -> None:
#         """Add an image to the document."""
#         try:
#             doc.add_picture(image_path, width=Inches(5.8))
#             # Center the last paragraph (the image)
#             doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
#             doc.add_paragraph()
#         except Exception:
#             doc.add_paragraph(f"(Failed to load image: {image_path})")

"""
Word document generation module.
"""

import re
from pathlib import Path
from typing import Optional, Union
import shutil

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from .executor import ExecutionResult

DEFAULT_ROLL_NO = "123102115"
_INVALID_XML_CHARS = re.compile(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]')

def _sanitize_text(text: str) -> str:
    return _INVALID_XML_CHARS.sub('', text)

class DocumentGenerator:
    """Generates Word documents from script execution results."""
    
    MONOSPACE_FONT = "Courier New"
    CODE_FONT_SIZE = Pt(10)
    
    def __init__(self, include_source: bool = False, verbose: bool = False, roll_no: Optional[str] = None):
        self.include_source = include_source
        self.verbose = verbose
        self.roll_no = roll_no if roll_no else DEFAULT_ROLL_NO
    
    def generate(self, result: ExecutionResult, output_path: Optional[Union[str, Path]] = None) -> Path:
        if output_path:
            output_path = Path(output_path)
        else:
            filename = f"{self.roll_no}_{result.script_path.stem}.docx"
            output_path = result.script_path.parent / filename
        
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        if self.verbose:
            print(f"[INFO] Generating document: {output_path}")
        
        doc = Document()
        self._setup_styles(doc)
        
        self._add_header(doc, result)
        
        if self.include_source:
            self._add_source_code(doc, result)
        
        self._add_output_section(doc, result)

        doc.save(str(output_path))
        self._cleanup_images(result)
        
        if self.verbose:
            print(f"[INFO] Document saved successfully: {output_path}")
        
        return output_path
    
    def _cleanup_images(self, result: ExecutionResult) -> None:
        pass
    
    def _setup_styles(self, doc: Document) -> None:
        styles = doc.styles
        if "Code" not in [s.name for s in styles]:
            code_style = styles.add_style("Code", WD_STYLE_TYPE.PARAGRAPH)
            code_style.font.name = self.MONOSPACE_FONT
            code_style.font.size = self.CODE_FONT_SIZE
            code_style.paragraph_format.space_before = Pt(6)
            code_style.paragraph_format.space_after = Pt(6)
    
    def _add_header(self, doc: Document, result: ExecutionResult) -> None:
        title = doc.add_heading(result.script_path.stem, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if self.roll_no:
            roll_para = doc.add_paragraph(f"Roll No: {self.roll_no}")
            roll_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
    
    def _add_source_code(self, doc: Document, result: ExecutionResult) -> None:
        doc.add_heading("Source Code", level=1)
        if result.source_code and result.source_code.strip():
            self._add_code_block(doc, result.source_code)
        else:
            doc.add_paragraph("(Empty file)")
        doc.add_paragraph()
    
    def _add_output_section(self, doc: Document, result: ExecutionResult) -> None:
        doc.add_heading("Output", level=1)
        
        if result.output_items:
            self._add_output_items(doc, result.output_items)
        elif result.stdout.strip():
            self._add_code_block(doc, result.stdout)
        else:
            doc.add_paragraph("(No output)")
        doc.add_paragraph()
    
    def _add_code_block(self, doc: Document, code: str, is_error: bool = False) -> None:
        code = _sanitize_text(code)
        lines = code.split("\n")
        for line in lines:
            para = doc.add_paragraph()
            run = para.add_run(line if line else " ")
            run.font.name = self.MONOSPACE_FONT
            run.font.size = self.CODE_FONT_SIZE
            if is_error:
                run.font.color.rgb = RGBColor(180, 0, 0)
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.line_spacing = 1.0

    def _add_output_items(self, doc: Document, items: list) -> None:
        """Add ordered output items (text, markdown, and images)."""
        for item in items:
            kind = getattr(item, "kind", None)
            if kind == "image":
                self._add_image(doc, item.content)
            elif kind == "markdown":
                self._add_markdown_block(doc, item.content)
            else:
                text = item.content if hasattr(item, "content") else ""
                if text and text.strip(): # Only add non-empty text blocks
                    self._add_code_block(doc, text)

    def _add_markdown_block(self, doc: Document, content: str) -> None:
        """Add markdown content as formatted text/headings."""
        content = _sanitize_text(content)
        lines = content.strip().split("\n")
        for line in lines:
            if line.startswith("# "):
                doc.add_heading(_sanitize_text(line[2:]), level=1)
            elif line.startswith("## "):
                doc.add_heading(_sanitize_text(line[3:]), level=2)
            elif line.startswith("### "):
                doc.add_heading(_sanitize_text(line[4:]), level=3)
            elif line.strip():
                # Regular paragraph
                doc.add_paragraph(_sanitize_text(line))
        # No extra spacing needed inside markdown blocks generally

    def _add_image(self, doc: Document, image_path: str) -> None:
        try:
            doc.add_picture(image_path, width=Inches(5.8))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
        except Exception:
            doc.add_paragraph(f"(Failed to load image: {image_path})")